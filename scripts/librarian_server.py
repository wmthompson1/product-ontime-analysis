"""
scripts/librarian_server.py
Librarian MCP server for Manufacturing Resource Planning (MRP) documents.

Exposes tools for listing local MRP documents, reading .docx content in-memory,
and committing concept-to-schema research artifacts into ArangoDB.

Adapted for the Replit/Linux repo from the private Windows prototype:
  * Default document root -> <repo>/docs/my-mrp-kb                         (env: MRP_DOCUMENTS_DIR)
  * Default border root   -> <repo>/certificate_for_receiving/border_extracts (env: BORDER_EXTRACTS_DIR)
  * ArangoDB password falls back to ARANGO_ROOT_PASSWORD (this repo's secret name).
  * commit_to_arangodb stays GATED OFF (MRP_ENABLE_GRAPH_COMMIT, default false) and
    targets a SEPARATE research database (default `mrp_research`) so it can never
    touch the certified `manufacturing_graph`.
"""

from __future__ import annotations

import csv
import json
import logging
import os
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Union

from dotenv import load_dotenv
import sqlglot
from sqlglot import exp

try:
    from mcp.server.fastmcp import FastMCP
except ImportError as exc:  # pragma: no cover - dependency guard
    raise ImportError(
        "The 'mcp' package is required. Install repository dependencies first."
    ) from exc

load_dotenv()

logging.basicConfig(
    level=os.getenv("LOG_LEVEL", "INFO").upper(),
    format="%(asctime)s %(levelname)s %(name)s - %(message)s",
)
logger = logging.getLogger("librarian_server")

SERVER_NAME = "mrp-librarian"

_REPO_ROOT = Path(__file__).resolve().parent.parent
_SCRIPTS_DIR = Path(__file__).resolve().parent


def _discover_default_doc_root() -> Path:
    configured = os.getenv("MRP_DOCUMENTS_DIR")
    if configured:
        return Path(configured).expanduser()

    return _REPO_ROOT / "docs" / "my-mrp-kb"


DEFAULT_DOC_ROOT = _discover_default_doc_root()
DEFAULT_BORDER_EXTRACT_ROOT = Path(
    os.getenv(
        "BORDER_EXTRACTS_DIR",
        str(_REPO_ROOT / "certificate_for_receiving" / "border_extracts"),
    )
).expanduser()
ARANGO_HOST = os.getenv("ARANGO_HOST", "http://127.0.0.1:8529")
# Research graph lives in its OWN database, separate from the certified
# `manufacturing_graph`. We deliberately do NOT read the shared `ARANGO_DB`
# here (in this repo that points at the canonical database) so research writes
# can never contaminate the approved semantic layer. Override only via the
# dedicated MRP_RESEARCH_ARANGO_DB / ARANGO_RESEARCH_DB knobs.
ARANGO_DATABASE = os.getenv(
    "MRP_RESEARCH_ARANGO_DB", os.getenv("ARANGO_RESEARCH_DB", "mrp_research")
)
ARANGO_USER = os.getenv("ARANGO_USER", "root")
ARANGO_PASSWORD = os.getenv("ARANGO_PASSWORD", os.getenv("ARANGO_ROOT_PASSWORD", ""))
NODE_COLLECTION = os.getenv("ARANGO_NODE_COLLECTION", "ai_research_node")
EDGE_COLLECTION = os.getenv("ARANGO_EDGE_COLLECTION", "ai_research_edge")
GRAPH_COMMIT_ENABLED = os.getenv("MRP_ENABLE_GRAPH_COMMIT", "false").strip().lower() in {
    "1",
    "true",
    "yes",
    "on",
}

# Certified graph identities that research writes must NEVER target.
CERTIFIED_DB_NAMES = {"manufacturing_graph", "semantic_graph"}
# Research collections live under this namespace so they can never collide with
# certified vertex/edge/bridge collections.
RESEARCH_COLLECTION_PREFIX = "ai_research"

mcp = FastMCP(SERVER_NAME)


def _normalize_arango_host(raw_host: Optional[str]) -> str:
    raw = (raw_host or "http://127.0.0.1:8529").strip()
    if "arangodb.cloud" in raw and ":" not in raw.split("//", 1)[-1]:
        raw = f"{raw}:8529"
    return raw


def _graph_commit_enabled() -> bool:
    return os.getenv("MRP_ENABLE_GRAPH_COMMIT", "false").strip().lower() in {
        "1",
        "true",
        "yes",
        "on",
    }


def _research_targets() -> Dict[str, str]:
    """Resolve and hard-guard the research write target.

    This is the single isolation chokepoint: it refuses any database that resolves
    to the certified graph and any collection outside the research namespace, so no
    caller (or stray env override) can contaminate the SME-approved layer.
    """
    db_name = os.getenv(
        "MRP_RESEARCH_ARANGO_DB", os.getenv("ARANGO_RESEARCH_DB", "mrp_research")
    ).strip()
    certified_db = (os.getenv("ARANGO_DB") or "").strip()
    forbidden = {name for name in (certified_db, *CERTIFIED_DB_NAMES) if name}
    if db_name in forbidden:
        raise RuntimeError(
            f"Refusing to commit: research database '{db_name}' resolves to a "
            f"certified database ({sorted(forbidden)}). Research writes must target a "
            "separate database (default 'mrp_research')."
        )

    node_collection = os.getenv("ARANGO_NODE_COLLECTION", "ai_research_node").strip()
    edge_collection = os.getenv("ARANGO_EDGE_COLLECTION", "ai_research_edge").strip()
    for collection in (node_collection, edge_collection):
        if not collection.startswith(RESEARCH_COLLECTION_PREFIX):
            raise RuntimeError(
                f"Refusing to commit: research collection '{collection}' must be "
                f"namespaced under '{RESEARCH_COLLECTION_PREFIX}' to stay isolated "
                "from certified collections."
            )

    return {
        "host": _normalize_arango_host(os.getenv("ARANGO_HOST", ARANGO_HOST)),
        "database": db_name,
        "user": os.getenv("ARANGO_USER", ARANGO_USER),
        "password": os.getenv("ARANGO_PASSWORD", os.getenv("ARANGO_ROOT_PASSWORD", "")),
        "node_collection": node_collection,
        "edge_collection": edge_collection,
    }


def _is_within_root(candidate: Path, root: Path) -> bool:
    try:
        candidate.resolve().relative_to(root.resolve())
        return True
    except ValueError:
        return False


def _resolve_path(path_text: Optional[str], allow_directory: bool = False) -> Path:
    base_root = DEFAULT_DOC_ROOT
    raw = Path(path_text) if path_text else base_root
    candidate = raw if raw.is_absolute() else base_root / raw
    candidate = candidate.resolve()

    if not base_root.exists():
        raise FileNotFoundError(f"Document root not found: {base_root}")
    if not _is_within_root(candidate, base_root):
        raise ValueError(f"Path escapes the configured document root: {candidate}")
    if not candidate.exists():
        raise FileNotFoundError(f"Path not found: {candidate}")
    if not allow_directory and not candidate.is_file():
        raise ValueError(f"Expected a file path, received directory: {candidate}")
    return candidate


def _resolve_border_extract_path(path_text: str) -> Path:
    base_root = DEFAULT_BORDER_EXTRACT_ROOT
    raw = Path(path_text)
    candidate = raw if raw.is_absolute() else base_root / raw
    candidate = candidate.resolve()

    if not base_root.exists():
        raise FileNotFoundError(f"Border extract root not found: {base_root}")
    if not _is_within_root(candidate, base_root):
        raise ValueError(f"Path escapes the configured border extract root: {candidate}")
    if not candidate.exists():
        raise FileNotFoundError(f"Path not found: {candidate}")
    if not candidate.is_file():
        raise ValueError(f"Expected a file path, received directory: {candidate}")
    return candidate


def _list_files(directory: Path) -> List[str]:
    files: List[str] = []
    for item in sorted(directory.rglob("*")):
        if item.is_file():
            files.append(str(item.relative_to(directory)).replace("\\", "/"))
    return files


def _extract_docx_paragraphs(document: Any) -> Iterable[str]:
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            yield text

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                yield " | ".join(cells)


def _extract_view_dependencies(view_name: str, raw_ddl: str) -> List[Dict[str, Any]]:
    lineage_edges: List[Dict[str, Any]] = []
    statements = sqlglot.parse(raw_ddl, read="tsql", error_level=sqlglot.ErrorLevel.IGNORE)

    for statement in statements:
        if not statement:
            continue

        for table in statement.find_all(exp.Table):
            src_table = table.sql(dialect="tsql").upper()
            if src_table.startswith("#") or src_table.startswith("@") or src_table == view_name:
                continue
            lineage_edges.append(
                {
                    "_from": f"physical_schema/{src_table}",
                    "_to": f"physical_schema/{view_name}",
                    "predicate": "DEPENDS_ON",
                    "context_hint": "Derived via SQL Server View.",
                }
            )

    return lineage_edges


def _normalize_handle(value: str, node_collection: str) -> str:
    if "/" in value:
        collection = value.split("/", 1)[0]
        if collection != node_collection:
            raise ValueError(
                f"Edge handle '{value}' references collection '{collection}', which is "
                f"not the guarded research node collection '{node_collection}'. Refusing "
                "to build an edge that could point outside the research graph."
            )
        return value
    return f"{node_collection}/{value}"


def _coerce_payload(payload: Union[str, Dict[str, Any]]) -> Dict[str, Any]:
    if isinstance(payload, str):
        try:
            payload = json.loads(payload)
        except json.JSONDecodeError as exc:
            raise ValueError("payload must be valid JSON or a JSON object") from exc
    if not isinstance(payload, dict):
        raise ValueError("payload must be a JSON object")
    return payload


def _node_key(node: Dict[str, Any]) -> str:
    node_type = str(node.get("node_type", "")).strip().lower()
    if node_type == "physical_schema":
        table_name = node.get("table_name")
        field_name = node.get("field_name")
        family = node.get("family")
        perspective = node.get("perspective")
        if not all([table_name, field_name, family, perspective]):
            raise ValueError(
                "Physical schema nodes must include table_name, field_name, family, and perspective"
            )
        return f"{table_name}:{field_name}:{family}:{perspective}"

    for field in ("_key", "key", "id", "name", "concept_key"):
        value = node.get(field)
        if value:
            return str(value)
    raise ValueError("Each node must include a key field such as _key, key, id, or name")


def _node_role(node: Optional[Dict[str, Any]]) -> str:
    if not node:
        return ""
    for field in ("node_type", "kind", "category", "entity_type", "role", "type"):
        value = node.get(field)
        if not value:
            continue
        role = str(value).strip().lower()
        if role in {"concept", "logical_concept", "semantic_concept"}:
            return "concept"
        if role in {"physical_schema", "schema", "table", "column", "database"}:
            return "physical_schema"
    return ""


def _build_node_docs(nodes: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    docs: List[Dict[str, Any]] = []
    for node in nodes:
        if not isinstance(node, dict):
            raise ValueError("Each node must be a JSON object")
        doc = dict(node)
        doc["_key"] = _node_key(node)
        if str(doc.get("node_type", "")).strip().lower() == "physical_schema":
            doc.pop("unique_id", None)
        docs.append(doc)
    return docs


def _build_edge_docs(
    edges: List[Dict[str, Any]],
    nodes_by_key: Dict[str, Dict[str, Any]],
    node_collection: str,
) -> List[Dict[str, Any]]:
    docs: List[Dict[str, Any]] = []
    for edge in edges:
        if not isinstance(edge, dict):
            raise ValueError("Each edge must be a JSON object")

        raw_from = edge.get("_from") or edge.get("from") or edge.get("source")
        raw_to = edge.get("_to") or edge.get("to") or edge.get("target")
        if not raw_from or not raw_to:
            raise ValueError("Each edge must include source and target handles")

        source_handle = _normalize_handle(str(raw_from), node_collection)
        target_handle = _normalize_handle(str(raw_to), node_collection)
        source_key = source_handle.split("/", 1)[-1]
        target_key = target_handle.split("/", 1)[-1]

        source_node = nodes_by_key.get(source_key)
        target_node = nodes_by_key.get(target_key)
        source_role = _node_role(source_node)
        target_role = _node_role(target_node)

        predicate = edge.get("predicate") or edge.get("label") or edge.get("relationship")
        if source_role == "concept" and target_role == "physical_schema":
            predicate = "RESOLVES_TO"
        elif not predicate:
            predicate = "RELATED_TO"

        doc = dict(edge)
        doc["_from"] = source_handle
        doc["_to"] = target_handle
        doc["predicate"] = predicate
        docs.append(doc)
    return docs


@mcp.tool()
def list_mrp_documents(directory: Optional[str] = None) -> Dict[str, Any]:
    """List files in the configured MRP document root or a subdirectory."""
    try:
        target = _resolve_path(directory, allow_directory=True)
        if not target.is_dir():
            raise ValueError(f"Expected a directory path, received file: {target}")
        files = _list_files(target)
        logger.info("Listed %d document(s) in %s", len(files), target)
        return {"directory": str(target), "count": len(files), "files": files}
    except Exception:
        logger.exception("Failed to list MRP documents")
        raise


@mcp.tool()
def read_document(filepath: str) -> str:
    """Read a .docx document from disk and return its text content."""
    try:
        target = _resolve_path(filepath)
        if target.suffix.lower() != ".docx":
            raise ValueError("read_document only supports .docx files")

        try:
            from docx import Document  # type: ignore
        except ImportError as exc:  # pragma: no cover - dependency guard
            raise ImportError(
                "The 'python-docx' package is required to read .docx files."
            ) from exc

        with target.open("rb") as handle:
            document = Document(BytesIO(handle.read()))

        text = "\n".join(_extract_docx_paragraphs(document))
        logger.info("Read %s (%d characters)", target, len(text))
        return text
    except Exception:
        logger.exception("Failed to read document")
        raise


@mcp.tool()
def commit_to_arangodb(payload: Union[str, Dict[str, Any]]) -> Dict[str, Any]:
    """Upsert research nodes and edges into the isolated mrp_research ArangoDB graph."""
    try:
        if not _graph_commit_enabled():
            raise RuntimeError(
                "Graph commit is disabled. Set MRP_ENABLE_GRAPH_COMMIT=true to enable commit_to_arangodb."
            )

        # Resolve + hard-guard the target so this can never touch the certified graph.
        targets = _research_targets()
        database = targets["database"]
        node_collection = targets["node_collection"]
        edge_collection = targets["edge_collection"]

        body = _coerce_payload(payload)
        nodes = body.get("nodes", [])
        edges = body.get("edges", [])

        if not isinstance(nodes, list) or not isinstance(edges, list):
            raise ValueError("payload must contain 'nodes' and 'edges' arrays")

        node_docs = _build_node_docs(nodes)
        nodes_by_key = {doc["_key"]: doc for doc in node_docs}
        edge_docs = _build_edge_docs(edges, nodes_by_key, node_collection)

        try:
            from arango import ArangoClient  # type: ignore
        except ImportError as exc:  # pragma: no cover - dependency guard
            raise ImportError(
                "The 'python-arango' package is required to commit to ArangoDB."
            ) from exc

        client = ArangoClient(hosts=targets["host"])
        system_db = client.db("_system", username=targets["user"], password=targets["password"])
        if not system_db.has_database(database):
            logger.info("Creating ArangoDB database %s", database)
            system_db.create_database(database)

        db = client.db(database, username=targets["user"], password=targets["password"])
        if not db.has_collection(node_collection):
            db.create_collection(node_collection)
        if not db.has_collection(edge_collection):
            db.create_collection(edge_collection, edge=True)

        node_upsert_query = f"""
        FOR node IN @nodes
          UPSERT {{ _key: node._key }}
          INSERT node
          UPDATE MERGE(OLD, node)
          IN {node_collection}
          OPTIONS {{ keepNull: false }}
          RETURN NEW
        """

        edge_upsert_query = f"""
        FOR edge IN @edges
          UPSERT {{ _from: edge._from, _to: edge._to, predicate: edge.predicate }}
          INSERT edge
          UPDATE MERGE(OLD, edge)
          IN {edge_collection}
          OPTIONS {{ keepNull: false }}
          RETURN NEW
        """

        node_result = list(db.aql.execute(node_upsert_query, bind_vars={"nodes": node_docs}))
        edge_result = list(db.aql.execute(edge_upsert_query, bind_vars={"edges": edge_docs}))

        logger.info(
            "Upserted %d node(s) and %d edge(s) into ArangoDB database %s",
            len(node_result),
            len(edge_result),
            database,
        )
        return {
            "database": database,
            "node_collection": node_collection,
            "edge_collection": edge_collection,
            "nodes_upserted": len(node_result),
            "edges_upserted": len(edge_result),
        }
    except Exception:
        logger.exception("Failed to commit payload to ArangoDB")
        raise


@mcp.tool()
def parse_ddl_csv(csv_filepath: str) -> str:
    """
    Read a CSV of view definitions and extract physical dependencies from view_ddl.
    Returns a JSON payload suitable for ArangoDB lineage ingestion.
    """
    try:
        target = _resolve_border_extract_path(csv_filepath)
        logger.info("Parsing view lineage from CSV: %s", target)

        lineage_edges: List[Dict[str, Any]] = []
        with target.open(mode="r", encoding="utf-8", newline="") as file:
            csv_reader = csv.DictReader(file)
            for row in csv_reader:
                schema_name = (row.get("schema_name") or "").strip()
                view_name = (row.get("view_name") or "").strip()
                raw_ddl = (row.get("view_ddl") or "").strip()
                if not schema_name or not view_name or not raw_ddl:
                    continue

                full_view_name = f"{schema_name}.{view_name}".upper()
                lineage_edges.extend(_extract_view_dependencies(full_view_name, raw_ddl))

        payload = {"operation": "upsert_lineage", "edges": lineage_edges}
        logger.info("Extracted %d lineage edge(s) from %s", len(lineage_edges), target)
        return json.dumps(payload, indent=2)
    except Exception:
        logger.exception("Failed to parse DDL CSV")
        raise


@mcp.tool()
def stage_terminology(filepath: Optional[str] = None, commit: bool = False) -> Dict[str, Any]:
    """Read a manufacturing terminology .docx and stage PROPOSED terms into mrp_research.

    Deterministic glossary extraction anchors each proposed term to the existing
    SME-approved perspectives/categories. The result is staged as PROPOSALS into the
    SEPARATE `mrp_research` graph only — the certified `manufacturing_graph` is never
    touched. By default this is a dry run that writes a reviewable JSON+CSV artifact;
    set commit=True to also upsert the proposals into the research graph.

    Args:
        filepath: Optional terminology document path (relative to the librarian
            document root). Defaults to the bundled terminology document.
        commit: When True, stage the proposals into the mrp_research graph via the
            gated, hard-guarded commit path. When False (default), dry run only.
    """
    try:
        if str(_SCRIPTS_DIR) not in sys.path:
            sys.path.insert(0, str(_SCRIPTS_DIR))
        import mrp_terminology_stager as stager

        doc_path = _resolve_path(filepath) if filepath else stager.DEFAULT_DOC
        logger.info("Staging terminology from %s (commit=%s)", doc_path, commit)

        return stager.run(
            doc_path=Path(doc_path),
            db_path=stager.DEFAULT_SQLITE,
            staging_root=stager.DEFAULT_STAGING_ROOT,
            commit=commit,
        )
    except Exception:
        logger.exception("Failed to stage terminology")
        raise


if __name__ == "__main__":
    mcp.run()
