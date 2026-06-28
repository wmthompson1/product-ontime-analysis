"""
scripts/my_mrp_ingestion_runner.py
Deterministic ingestion runner for the My MRP learning-loop corpus.

This script processes files in the configured knowledge-base folder, emits
versioned evaluation artifacts, and records human-auditable reasoning traces. It
does not mutate SQL schema and does not write to ArangoDB directly.

Adapted for the Replit/Linux repo from the private Windows prototype:
  * Default watch root  -> <repo>/docs/my-mrp-kb        (env: MRP_DOCUMENTS_DIR)
  * Default artifacts   -> <repo>/mrp_ingest_artifacts  (env: MRP_INGEST_ARTIFACT_ROOT)
The artifact root is intentionally OUTSIDE the watch root so the runner never
ingests its own output on a later cycle.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import logging
import os
import re
import time
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List, Optional, Sequence, Tuple

from dotenv import load_dotenv

load_dotenv()

logging.basicConfig(
    level=os.getenv("LOG_LEVEL", "INFO").upper(),
    format="%(asctime)s %(levelname)s %(name)s - %(message)s",
)
logger = logging.getLogger("my_mrp_ingestion_runner")

_REPO_ROOT = Path(__file__).resolve().parent.parent

DEFAULT_WATCH_ROOT = Path(
    os.getenv("MRP_DOCUMENTS_DIR", str(_REPO_ROOT / "docs" / "my-mrp-kb"))
).expanduser()

DEFAULT_ARTIFACT_ROOT = Path(
    os.getenv("MRP_INGEST_ARTIFACT_ROOT", str(_REPO_ROOT / "mrp_ingest_artifacts"))
).expanduser()

TEXT_EXTENSIONS = {".md", ".txt", ".docx"}
MEDIA_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".svg"}

PII_PATTERNS: Dict[str, re.Pattern[str]] = {
    "email": re.compile(r"\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b"),
    "ssn": re.compile(r"\b\d{3}-\d{2}-\d{4}\b"),
    "phone": re.compile(r"\b(?:\+?1[\s.-]?)?(?:\(?\d{3}\)?[\s.-]?)\d{3}[\s.-]?\d{4}\b"),
}


@dataclass
class DocumentDecision:
    relative_path: str
    extension: str
    status: str
    sha256: str
    reason: str
    pii_hits: Dict[str, int]
    last_modified_utc: str


def _now_utc() -> datetime:
    return datetime.now(timezone.utc)


def _format_ts(ts: datetime) -> str:
    return ts.isoformat().replace("+00:00", "Z")


def _is_within_root(candidate: Path, root: Path) -> bool:
    try:
        candidate.resolve().relative_to(root.resolve())
        return True
    except ValueError:
        return False


def _sha256_bytes(raw: bytes) -> str:
    return hashlib.sha256(raw).hexdigest()


def _sha256_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8", errors="ignore")).hexdigest()


def _read_docx_text(path: Path) -> str:
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:  # pragma: no cover
        raise ImportError("python-docx is required to process .docx files") from exc

    doc = Document(str(path))
    chunks: List[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                chunks.append(row_text)

    return "\n".join(chunks)


def _read_text(path: Path) -> str:
    if path.suffix.lower() == ".docx":
        return _read_docx_text(path)

    return path.read_text(encoding="utf-8", errors="ignore")


def _detect_pii(content: str) -> Dict[str, int]:
    hits: Dict[str, int] = {}
    if not content:
        return hits

    for category, pattern in PII_PATTERNS.items():
        matches = pattern.findall(content)
        if matches:
            hits[category] = len(matches)
    return hits


def _collect_files(root: Path) -> List[Path]:
    files = [path for path in root.rglob("*") if path.is_file()]
    files.sort(key=lambda item: str(item.relative_to(root)).replace("\\", "/").lower())
    return files


def _build_corpus_fingerprint(file_hashes: Sequence[Tuple[str, str]]) -> str:
    stable = "\n".join(f"{rel}:{digest}" for rel, digest in file_hashes)
    return _sha256_text(stable)


def run_ingestion_once(watch_root: Path, artifact_root: Path) -> Dict[str, object]:
    if not watch_root.exists():
        raise FileNotFoundError(f"Watch root not found: {watch_root}")
    if not watch_root.is_dir():
        raise ValueError(f"Watch root must be a directory: {watch_root}")
    # The artifact root MUST stay outside the watch root, otherwise rglob() would
    # pick up our own run artifacts and self-ingest them on the next cycle. The
    # defaults already guarantee this; this guard catches a bad env/CLI override.
    if _is_within_root(artifact_root, watch_root):
        raise ValueError(
            "Artifact root must live OUTSIDE the watch root to avoid self-ingestion: "
            f"artifact_root={artifact_root} is inside watch_root={watch_root}"
        )

    run_start = _now_utc()
    run_id = run_start.strftime("%Y%m%dT%H%M%SZ")

    decisions: List[DocumentDecision] = []
    file_hashes: List[Tuple[str, str]] = []

    for path in _collect_files(watch_root):
        if not _is_within_root(path, watch_root):
            logger.warning("Skipped path outside root boundary: %s", path)
            continue

        rel = str(path.relative_to(watch_root)).replace("\\", "/")
        raw = path.read_bytes()
        sha = _sha256_bytes(raw)
        file_hashes.append((rel, sha))
        ext = path.suffix.lower()
        mtime = _format_ts(datetime.fromtimestamp(path.stat().st_mtime, tz=timezone.utc))

        if ext in TEXT_EXTENSIONS:
            content = _read_text(path)
            pii_hits = _detect_pii(content)
            if pii_hits:
                decisions.append(
                    DocumentDecision(
                        relative_path=rel,
                        extension=ext,
                        status="blocked_pii",
                        sha256=sha,
                        reason="PII pattern match detected; content excluded from ingestion.",
                        pii_hits=pii_hits,
                        last_modified_utc=mtime,
                    )
                )
            else:
                decisions.append(
                    DocumentDecision(
                        relative_path=rel,
                        extension=ext,
                        status="ingested",
                        sha256=sha,
                        reason="Text content approved within ingestion policy boundary.",
                        pii_hits={},
                        last_modified_utc=mtime,
                    )
                )
        elif ext in MEDIA_EXTENSIONS:
            decisions.append(
                DocumentDecision(
                    relative_path=rel,
                    extension=ext,
                    status="metadata_only",
                    sha256=sha,
                    reason="Media asset tracked as lineage metadata only.",
                    pii_hits={},
                    last_modified_utc=mtime,
                )
            )
        else:
            decisions.append(
                DocumentDecision(
                    relative_path=rel,
                    extension=ext,
                    status="skipped_unsupported",
                    sha256=sha,
                    reason="Unsupported extension for ingestion; retained in manifest.",
                    pii_hits={},
                    last_modified_utc=mtime,
                )
            )

    corpus_fingerprint = _build_corpus_fingerprint(file_hashes)
    run_dir = artifact_root / run_id
    run_dir.mkdir(parents=True, exist_ok=True)

    manifest = {
        "run_id": run_id,
        "run_started_utc": _format_ts(run_start),
        "watch_root": str(watch_root),
        "corpus_fingerprint": corpus_fingerprint,
        "file_count": len(file_hashes),
        "files": [{"relative_path": rel, "sha256": digest} for rel, digest in file_hashes],
    }

    reasoning_trace = {
        "run_id": run_id,
        "reasoning_model": "deterministic-rule-engine",
        "policy_decisions": [decision.__dict__ for decision in decisions],
    }

    ingested = sum(1 for d in decisions if d.status == "ingested")
    blocked = sum(1 for d in decisions if d.status == "blocked_pii")
    metadata_only = sum(1 for d in decisions if d.status == "metadata_only")

    evaluation = {
        "run_id": run_id,
        "summary": {
            "ingested_documents": ingested,
            "blocked_for_pii": blocked,
            "metadata_only_assets": metadata_only,
            "total_documents_seen": len(decisions),
        },
        "controls": {
            "deterministic_improvement_paths": True,
            "strict_data_boundary_compliance": True,
            "human_auditable_reasoning_traces": True,
            "versioned_evaluation_artifacts": True,
            "schema_mutation_attempted": False,
            "pii_ingestion_permitted": False,
            "uncontrolled_model_updates_permitted": False,
        },
        "notes": [
            "This runner does not execute database schema mutations.",
            "ArangoDB writes are intentionally excluded from this ingestion stage.",
            "Use explicit reviewed payloads before calling commit_to_arangodb.",
        ],
    }

    (run_dir / "manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    (run_dir / "reasoning_trace.json").write_text(
        json.dumps(reasoning_trace, indent=2),
        encoding="utf-8",
    )
    (run_dir / "evaluation.json").write_text(json.dumps(evaluation, indent=2), encoding="utf-8")

    logger.info(
        "Ingestion run %s complete: %d ingested, %d blocked, %d metadata-only",
        run_id,
        ingested,
        blocked,
        metadata_only,
    )

    return {
        "run_id": run_id,
        "artifact_dir": str(run_dir),
        "ingested_documents": ingested,
        "blocked_for_pii": blocked,
        "metadata_only_assets": metadata_only,
        "total_documents_seen": len(decisions),
    }


def watch_ingestion_loop(
    watch_root: Path,
    artifact_root: Path,
    interval_seconds: int,
    max_cycles: Optional[int],
) -> None:
    cycle = 0
    last_fingerprint: Optional[str] = None

    while True:
        cycle += 1
        result = run_ingestion_once(watch_root, artifact_root)

        current_run_dir = Path(result["artifact_dir"])
        manifest = json.loads((current_run_dir / "manifest.json").read_text(encoding="utf-8"))
        current_fingerprint = str(manifest["corpus_fingerprint"])

        if current_fingerprint == last_fingerprint:
            logger.info("No corpus change detected after cycle %d", cycle)
        else:
            logger.info("Corpus fingerprint updated on cycle %d", cycle)
            last_fingerprint = current_fingerprint

        if max_cycles is not None and cycle >= max_cycles:
            logger.info("Stopping watch loop after %d cycles", cycle)
            return

        time.sleep(interval_seconds)


def _parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Deterministic My MRP ingestion runner")
    parser.add_argument(
        "--watch-root",
        default=str(DEFAULT_WATCH_ROOT),
        help="Folder to ingest (default: docs/my-mrp-kb)",
    )
    parser.add_argument(
        "--artifact-root",
        default=str(DEFAULT_ARTIFACT_ROOT),
        help="Folder to store versioned ingestion artifacts",
    )
    parser.add_argument(
        "--watch",
        action="store_true",
        help="Run continuously in watch mode",
    )
    parser.add_argument(
        "--interval-seconds",
        type=int,
        default=30,
        help="Polling interval for --watch mode",
    )
    parser.add_argument(
        "--max-cycles",
        type=int,
        default=None,
        help="Optional max number of watch cycles",
    )
    return parser.parse_args()


def main() -> None:
    args = _parse_args()
    watch_root = Path(args.watch_root).expanduser()
    artifact_root = Path(args.artifact_root).expanduser()

    try:
        if args.watch:
            watch_ingestion_loop(
                watch_root=watch_root,
                artifact_root=artifact_root,
                interval_seconds=args.interval_seconds,
                max_cycles=args.max_cycles,
            )
        else:
            result = run_ingestion_once(watch_root=watch_root, artifact_root=artifact_root)
            print(json.dumps(result, indent=2))
    except Exception:
        logger.exception("My MRP ingestion run failed")
        raise


if __name__ == "__main__":
    main()
