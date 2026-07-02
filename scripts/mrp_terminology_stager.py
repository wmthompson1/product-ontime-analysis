"""
scripts/mrp_terminology_stager.py
Deterministic terminology stager for the "My MRP" librarian.

Reads a manufacturing/MRP terminology .docx from the librarian document root,
extracts candidate terms (deterministically from the Glossary by default),
anchors each candidate to the EXISTING SME-approved perspectives and categories
(read-only from the certified SQLite metadata), and stages the result as
PROPOSALS into the SEPARATE `mrp_research` ArangoDB graph.

Governance (the "Solder Pattern"):
  * Perspective/concept DEFINITIONS are SME-approved ground truth. Document-derived
    terms here are PROPOSALS ONLY. They are tagged `approval_status="proposed"`,
    `source_type="document_extraction"`, `certified=false`, and never become
    approved definitions automatically.
  * Approved perspectives/categories are COPIED into the research graph as
    `approved_anchor` reference nodes; candidate edges point at those research-side
    anchors, so NO edge ever references a certified-graph handle.
  * Live writes go ONLY through the librarian's gated `commit_to_arangodb`, which
    refuses any target that resolves to the certified database/collections.
  * Default run is a DRY RUN: it writes a reviewable JSON+CSV artifact and touches
    no database/network. Live staging requires an explicit `--commit`.
"""

from __future__ import annotations

import argparse
import csv
import json
import logging
import os
import re
import sys
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Set, Tuple

from dotenv import load_dotenv

load_dotenv()

logging.basicConfig(
    level=os.getenv("LOG_LEVEL", "INFO").upper(),
    format="%(asctime)s %(levelname)s %(name)s - %(message)s",
)
logger = logging.getLogger("mrp_terminology_stager")

_REPO_ROOT = Path(__file__).resolve().parent.parent
_SCRIPTS_DIR = Path(__file__).resolve().parent
if str(_SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(_SCRIPTS_DIR))

# Reuse the ingestion runner's deterministic safety helpers (PII scan + hashing).
from my_mrp_ingestion_runner import _detect_pii, _sha256_bytes  # noqa: E402

DEFAULT_DOC_ROOT = Path(
    os.getenv("MRP_DOCUMENTS_DIR", str(_REPO_ROOT / "docs" / "my-mrp-kb"))
).expanduser()
DEFAULT_DOC = DEFAULT_DOC_ROOT / "Manufacturing and MRP Terminology in Semantic Models.docx"
DEFAULT_SQLITE = (
    _REPO_ROOT / "hf-space-inventory-sqlgen" / "app_schema" / "manufacturing.db"
)
DEFAULT_STAGING_ROOT = Path(
    os.getenv("MRP_RESEARCH_STAGING_ROOT", str(_REPO_ROOT / "mrp_research_staging"))
).expanduser()

EM_DASH = "\u2014"  # — separates a glossary term from its definition
STAR = "\u2605"  # ★ marks a foundational term
GLOSSARY_MARKER = "glossary of"
ANCHOR_SCORE_THRESHOLD = 2  # minimum token-overlap weight to record an anchor
NAME_TOKEN_WEIGHT = 2  # term-name token matches count double vs definition tokens

# Generic words that would create false anchors if they were allowed to match.
STOPWORDS: Set[str] = {
    "a", "an", "and", "are", "as", "at", "be", "by", "for", "from", "in", "into",
    "is", "of", "on", "or", "the", "to", "with", "that", "this", "these", "those",
    "it", "its", "their", "such", "via", "per", "each", "all", "used", "use",
    "data", "system", "systems", "model", "models", "semantic", "value", "values",
    "level", "levels", "based", "set", "type", "types", "via", "within", "across",
    "between", "through", "over", "using", "one", "unit", "units", "item", "items",
    "process", "processes", "management", "manufacturing", "production", "product",
    "products", "operations", "operation", "ontology", "owl", "rdf", "graph",
    "class", "property", "properties", "individual", "individuals", "objectproperty",
    "objectproperties", "datatypeproperty", "standard", "reference", "structure",
    "real", "world", "modern", "factory", "enterprise", "planning",
}


@dataclass
class CandidateTerm:
    term: str
    slug: str
    acronym: str
    foundational: bool
    definition: str
    source_section: str
    perspective_anchors: List[Dict[str, Any]] = field(default_factory=list)
    category_anchors: List[Dict[str, Any]] = field(default_factory=list)

    @property
    def is_anchored(self) -> bool:
        return bool(self.perspective_anchors or self.category_anchors)


def _now_utc() -> datetime:
    return datetime.now(timezone.utc)


def _format_ts(ts: datetime) -> str:
    return ts.isoformat().replace("+00:00", "Z")


def slugify(text: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "_", text.strip().lower()).strip("_")
    return re.sub(r"_+", "_", slug)


def _read_docx(path: Path) -> Tuple[List[str], str]:
    """Return (paragraph_lines, full_text_including_tables) for a .docx file."""
    try:
        from docx import Document  # type: ignore
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise ImportError("python-docx is required to read .docx files") from exc

    document = Document(str(path))
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]

    table_lines: List[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                table_lines.append(" | ".join(cells))

    full_text = "\n".join(paragraphs + table_lines)
    return paragraphs, full_text


def parse_glossary(paragraphs: Sequence[str]) -> List[CandidateTerm]:
    """Extract `Term (ACRONYM) ★ — definition` entries from the Glossary section."""
    terms: List[CandidateTerm] = []
    seen: Set[str] = set()
    in_glossary = False
    sep = f" {EM_DASH} "

    for line in paragraphs:
        if not in_glossary:
            if line.lower().startswith(GLOSSARY_MARKER):
                in_glossary = True
            continue

        if sep not in line:
            continue

        left, definition = line.split(sep, 1)
        left = left.strip()
        definition = definition.strip()

        # A glossary head is short and is not a prose sentence.
        if len(left) > 80 or left.endswith(".") or not definition:
            continue

        foundational = STAR in left
        left = left.replace(STAR, "").strip()

        acronym = ""
        acro_match = re.search(r"\(([^)]+)\)\s*$", left)
        if acro_match:
            acronym = acro_match.group(1).strip()
            left = left[: acro_match.start()].strip()

        term = left.strip()
        if not term:
            continue

        slug = slugify(term)
        if not slug or slug in seen:
            continue
        seen.add(slug)

        terms.append(
            CandidateTerm(
                term=term,
                slug=slug,
                acronym=acronym,
                foundational=foundational,
                definition=definition,
                source_section="Glossary",
            )
        )

    return terms


def tokenize(text: str) -> Set[str]:
    raw = re.findall(r"[a-z0-9]+", (text or "").lower())
    return {tok for tok in raw if len(tok) > 2 and tok not in STOPWORDS}


def load_anchors(db_path: Path) -> Dict[str, List[Dict[str, Any]]]:
    """Read the SME-approved perspectives and categories (read-only)."""
    import sqlite3

    if not db_path.exists():
        raise FileNotFoundError(f"SQLite metadata database not found: {db_path}")

    conn = sqlite3.connect(f"file:{db_path}?mode=ro", uri=True)
    conn.row_factory = sqlite3.Row
    try:
        perspectives = [
            {
                "name": r["perspective_name"],
                "description": r["description"] or "",
                "stakeholder_role": r["stakeholder_role"] or "",
                "priority_focus": r["priority_focus"] or "",
            }
            for r in conn.execute(
                "SELECT perspective_name, description, stakeholder_role, priority_focus "
                "FROM schema_perspectives ORDER BY perspective_name"
            ).fetchall()
        ]
        categories = [
            {"name": r["category_name"], "description": r["description"] or ""}
            for r in conn.execute(
                "SELECT category_name, description FROM schema_entity_categories "
                "ORDER BY category_name"
            ).fetchall()
        ]
    finally:
        conn.close()

    for p in perspectives:
        p["_tokens"] = tokenize(
            " ".join(
                [p["name"].replace("_", " "), p["description"], p["stakeholder_role"], p["priority_focus"]]
            )
        )
    for c in categories:
        c["_tokens"] = tokenize(" ".join([c["name"].replace("_", " "), c["description"]]))

    return {"perspectives": perspectives, "categories": categories}


def _score(term: CandidateTerm, anchor_tokens: Set[str]) -> Tuple[int, List[str]]:
    name_tokens = tokenize(term.term)
    def_tokens = tokenize(term.definition)
    matched: List[str] = []
    score = 0
    for tok in sorted(anchor_tokens):
        if tok in name_tokens:
            score += NAME_TOKEN_WEIGHT
            matched.append(tok)
        elif tok in def_tokens:
            score += 1
            matched.append(tok)
    return score, matched


def _best_anchors(
    term: CandidateTerm, anchors: Sequence[Dict[str, Any]]
) -> List[Dict[str, Any]]:
    scored = []
    for anchor in anchors:
        score, matched = _score(term, anchor["_tokens"])
        if score >= ANCHOR_SCORE_THRESHOLD:
            scored.append(
                {
                    "name": anchor["name"],
                    "score": score,
                    "matched_tokens": matched[:12],
                }
            )
    scored.sort(key=lambda a: (-a["score"], a["name"]))
    if not scored:
        return []
    top = scored[0]["score"]
    # Keep the top match plus any ties (a term may legitimately span perspectives).
    return [a for a in scored if a["score"] == top]


def anchor_terms(
    terms: Sequence[CandidateTerm], anchors: Dict[str, List[Dict[str, Any]]]
) -> None:
    for term in terms:
        term.perspective_anchors = _best_anchors(term, anchors["perspectives"])
        term.category_anchors = _best_anchors(term, anchors["categories"])


def build_payload(
    terms: Sequence[CandidateTerm],
    anchors: Dict[str, List[Dict[str, Any]]],
    doc_rel_path: str,
    doc_sha256: str,
    extraction_method: str,
) -> Dict[str, List[Dict[str, Any]]]:
    nodes: List[Dict[str, Any]] = []
    edges: List[Dict[str, Any]] = []
    referenced_anchor_keys: Set[str] = set()

    anchor_meta: Dict[str, Dict[str, Any]] = {}
    for p in anchors["perspectives"]:
        anchor_meta[f"anchor_perspective__{slugify(p['name'])}"] = {
            "anchor_kind": "perspective",
            "name": p["name"],
            "description": p["description"],
        }
    for c in anchors["categories"]:
        anchor_meta[f"anchor_category__{slugify(c['name'])}"] = {
            "anchor_kind": "category",
            "name": c["name"],
            "description": c["description"],
        }

    for term in terms:
        term_key = f"term__{term.slug}"
        nodes.append(
            {
                "_key": term_key,
                "node_type": "proposed_term",
                "name": term.term,
                "term": term.term,
                "acronym": term.acronym,
                "foundational": term.foundational,
                "proposed_definition": term.definition,
                "document_definition": term.definition,
                "approval_status": "proposed",
                "source_type": "document_extraction",
                "certified": False,
                "source_doc": doc_rel_path,
                "doc_sha256": doc_sha256,
                "source_section": term.source_section,
                "extraction_method": extraction_method,
                "is_anchored": term.is_anchored,
                "reviewer_decision": "proposed",
            }
        )

        for kind, anchor_list, predicate, prefix in (
            ("perspective", term.perspective_anchors, "CANDIDATE_TERM_FOR_PERSPECTIVE", "anchor_perspective__"),
            ("category", term.category_anchors, "CANDIDATE_TERM_IN_CATEGORY", "anchor_category__"),
        ):
            for anchor in anchor_list:
                anchor_key = f"{prefix}{slugify(anchor['name'])}"
                referenced_anchor_keys.add(anchor_key)
                edges.append(
                    {
                        "_from": term_key,
                        "_to": anchor_key,
                        "predicate": predicate,
                        "match_score": anchor["score"],
                        "matched_tokens": ", ".join(anchor["matched_tokens"]),
                        "approval_status": "proposed",
                        "source_type": "document_extraction",
                    }
                )

    for anchor_key in sorted(referenced_anchor_keys):
        meta = anchor_meta[anchor_key]
        nodes.append(
            {
                "_key": anchor_key,
                "node_type": "approved_anchor",
                "anchor_kind": meta["anchor_kind"],
                "name": meta["name"],
                "description": meta["description"],
                "approval_status": "approved_reference",
                "source_type": "certified_metadata_reference",
                "certified": False,
                "mirrors_certified_name": meta["name"],
            }
        )

    return {"nodes": nodes, "edges": edges}


def write_artifacts(
    run_dir: Path,
    payload: Dict[str, List[Dict[str, Any]]],
    terms: Sequence[CandidateTerm],
    manifest: Dict[str, Any],
) -> None:
    run_dir.mkdir(parents=True, exist_ok=True)
    (run_dir / "proposed_terms.json").write_text(
        json.dumps(payload, indent=2), encoding="utf-8"
    )
    (run_dir / "manifest.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")

    with (run_dir / "proposed_terms.csv").open("w", encoding="utf-8", newline="") as fh:
        writer = csv.writer(fh)
        writer.writerow(
            [
                "term",
                "acronym",
                "foundational",
                "perspective_anchors",
                "category_anchors",
                "anchored",
                "definition",
                "reviewer_decision",
            ]
        )
        for term in terms:
            writer.writerow(
                [
                    term.term,
                    term.acronym,
                    "yes" if term.foundational else "",
                    "; ".join(f"{a['name']}({a['score']})" for a in term.perspective_anchors),
                    "; ".join(f"{a['name']}({a['score']})" for a in term.category_anchors),
                    "yes" if term.is_anchored else "no",
                    term.definition,
                    "proposed",
                ]
            )


def commit_payload(payload: Dict[str, List[Dict[str, Any]]]) -> Dict[str, Any]:
    """Stage the payload into the separate mrp_research graph via the gated librarian path.

    This does NOT enable the commit gate. Asking to commit is the intent; the
    environment gate is the control. Live writes require the operator to set
    MRP_ENABLE_GRAPH_COMMIT=true in the environment. The librarian's
    `commit_to_arangodb` enforces that gate plus the research-only hard guards
    (separate database + `ai_research`-namespaced collections).
    """
    if str(_SCRIPTS_DIR) not in sys.path:
        sys.path.insert(0, str(_SCRIPTS_DIR))
    import librarian_server  # noqa: E402

    return librarian_server.commit_to_arangodb(payload)


def run(
    doc_path: Path,
    db_path: Path,
    staging_root: Path,
    commit: bool = False,
    use_ai: bool = False,
) -> Dict[str, Any]:
    if not doc_path.exists():
        raise FileNotFoundError(f"Terminology document not found: {doc_path}")
    if use_ai:
        # AI drafting is intentionally not wired: document-derived content must
        # stay deterministic and auditable, and must never mint certified content.
        raise NotImplementedError(
            "AI extraction is intentionally disabled; deterministic glossary "
            "parsing is the governed default."
        )

    run_start = _now_utc()
    run_id = run_start.strftime("%Y%m%dT%H%M%SZ")

    try:
        doc_rel_path = str(doc_path.resolve().relative_to(_REPO_ROOT))
    except ValueError:
        doc_rel_path = str(doc_path)

    doc_bytes = doc_path.read_bytes()
    doc_hash = _sha256_bytes(doc_bytes)

    paragraphs, full_text = _read_docx(doc_path)

    pii_hits = _detect_pii(full_text)
    if pii_hits:
        raise RuntimeError(
            f"PII detected in {doc_rel_path} ({pii_hits}); refusing to stage terminology."
        )

    terms = parse_glossary(paragraphs)
    if not terms:
        raise ValueError(
            f"No glossary terms parsed from {doc_rel_path}; expected a "
            "'Glossary of ...' section with 'Term — definition' entries."
        )

    anchors = load_anchors(db_path)
    anchor_terms(terms, anchors)

    extraction_method = "deterministic_glossary"
    payload = build_payload(terms, anchors, doc_rel_path, doc_hash, extraction_method)

    anchored = sum(1 for t in terms if t.is_anchored)
    manifest = {
        "run_id": run_id,
        "run_started_utc": _format_ts(run_start),
        "source_doc": doc_rel_path,
        "doc_sha256": doc_hash,
        "extraction_method": extraction_method,
        "pii_status": "clean",
        "terms_extracted": len(terms),
        "terms_anchored": anchored,
        "terms_unanchored": len(terms) - anchored,
        "proposed_term_nodes": sum(1 for n in payload["nodes"] if n["node_type"] == "proposed_term"),
        "approved_anchor_nodes": sum(1 for n in payload["nodes"] if n["node_type"] == "approved_anchor"),
        "candidate_edges": len(payload["edges"]),
        "committed": False,
        "target_research_db": os.getenv(
            "MRP_RESEARCH_ARANGO_DB", os.getenv("ARANGO_RESEARCH_DB", "mrp_research")
        ),
    }

    run_dir = staging_root / run_id

    commit_result: Optional[Dict[str, Any]] = None
    if commit:
        commit_result = commit_payload(payload)
        manifest["committed"] = True
        manifest["commit_result"] = commit_result

    write_artifacts(run_dir, payload, terms, manifest)

    summary = {
        "run_id": run_id,
        "artifact_dir": str(run_dir),
        "source_doc": doc_rel_path,
        "terms_extracted": len(terms),
        "terms_anchored": anchored,
        "terms_unanchored": len(terms) - anchored,
        "candidate_edges": len(payload["edges"]),
        "committed": bool(commit),
    }
    if commit_result is not None:
        summary["commit_result"] = commit_result

    logger.info(
        "Terminology stage %s: %d terms (%d anchored), %d edges, committed=%s",
        run_id,
        len(terms),
        anchored,
        len(payload["edges"]),
        bool(commit),
    )
    return summary


def _parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Deterministically stage MRP terminology proposals into mrp_research"
    )
    parser.add_argument("--doc", default=str(DEFAULT_DOC), help="Path to terminology .docx")
    parser.add_argument("--db", default=str(DEFAULT_SQLITE), help="Certified SQLite metadata DB")
    parser.add_argument(
        "--staging-root",
        default=str(DEFAULT_STAGING_ROOT),
        help="Folder for reviewable staging artifacts",
    )
    parser.add_argument(
        "--commit",
        action="store_true",
        help=(
            "Stage proposals into the separate mrp_research ArangoDB graph "
            "(also requires MRP_ENABLE_GRAPH_COMMIT=true in the environment)"
        ),
    )
    parser.add_argument(
        "--ai",
        action="store_true",
        help="(Disabled) AI extraction is intentionally not available",
    )
    return parser.parse_args()


def main() -> None:
    args = _parse_args()
    summary = run(
        doc_path=Path(args.doc).expanduser(),
        db_path=Path(args.db).expanduser(),
        staging_root=Path(args.staging_root).expanduser(),
        commit=args.commit,
        use_ai=args.ai,
    )
    print(json.dumps(summary, indent=2))


if __name__ == "__main__":
    main()
